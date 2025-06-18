import streamlit as st
import streamlit.components.v1 as components
from io import BytesIO
from docx import Document

# Pemetaan karakter ke Unicode Look-alike
unicode_map = {
    'a': 'а', 'e': 'е', 'i': 'і', 'o': 'о', 'p': 'р',
    'c': 'с', 'y': 'у', 'x': 'х', 'A': 'А', 'E': 'Е',
    'O': 'О', 'P': 'Р', 'C': 'С', 'Y': 'Ү', 'X': 'Х',
    'l': 'ӏ', 'B': 'В'
}

# Fungsi konversi
def to_lookalike(text):
    return ''.join(unicode_map.get(c, c) for c in text)

def convert_docx_text(doc):
    for para in doc.paragraphs:
        for run in para.runs:
            run.text = to_lookalike(run.text)
    if doc.tables:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.text = to_lookalike(run.text)
    return doc

# UI
st.set_page_config(page_title="y-tools", layout="centered")
st.title("y-tools")

# Navigasi antar halaman
page = st.sidebar.radio("📁 Pilih Halaman", ["Teks Manual", "File DOCX"])

if page == "Teks Manual":
    input_text = st.text_area("✏️ Masukkan Teks Anda", height=250, placeholder="Tulis teks panjang di sini...")

    if input_text:
        output_text = to_lookalike(input_text)

        st.markdown("### 🔍 Preview Hasil:")
        edited_text = st.text_area("🖥️ Hasil Unicode:", value=output_text, height=200, label_visibility="collapsed")

        components.html(f"""
            <script>
              function copyToClipboard(text) {{
                navigator.clipboard.writeText(text).then(function() {{
                  const msg = document.getElementById('copied-msg');
                  msg.style.display = 'inline';
                  setTimeout(() => msg.style.display = 'none', 1500);
                }});
              }}
            </script>
            <button onclick="copyToClipboard(`{edited_text}`)" style="
                background-color:#4CAF50;
                color:white;
                border:none;
                padding:10px 20px;
                margin-top:10px;
                border-radius:5px;
                cursor:pointer;
                font-weight:bold;
                transition: background-color 0.3s;">
                📋 Salin ke Clipboard
            </button>
            <span id="copied-msg" style="display:none; margin-left:10px; color:green; font-weight:bold;">
                ✅ Disalin!
            </span>
            """,
            height=70,
        )
    else:
        st.info("Masukkan teks di atas untuk melihat hasil konversinya.")

if page == "File DOCX":
    st.markdown("---")
    st.header("📄 Konversi File DOCX")

    uploaded_file = st.file_uploader("Unggah file .docx", type=["docx"])

    if uploaded_file:
        try:
            doc = Document(uploaded_file)
            converted_doc = convert_docx_text(doc)
            output = BytesIO()
            converted_doc.save(output)
            output.seek(0)
            filename = uploaded_file.name.encode('utf-8').decode('utf-8')
            st.download_button("⬇️ Unduh Hasil", output, file_name=f"converted_{filename}")
        except Exception as e:
            st.error(f"Terjadi kesalahan saat memproses file: {e}")
